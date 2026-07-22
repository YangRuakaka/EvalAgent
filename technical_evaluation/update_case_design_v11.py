from pathlib import Path

from docx import Document
from docx.shared import RGBColor


ROOT = Path(r"D:\Github\EvalAgent")
SOURCE = ROOT / "output" / "EvalAgent_WebHarbor_Case_Design.docx"
OUTPUT = ROOT / "output" / "EvalAgent_WebHarbor_Case_Design_v1.1.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    """Replace visible text while retaining the paragraph's first-run formatting."""
    if paragraph.runs:
        anchor = paragraph.runs[0]
        props = {
            "bold": anchor.bold,
            "italic": anchor.italic,
            "underline": anchor.underline,
            "font_name": anchor.font.name,
            "font_size": anchor.font.size,
            "font_color": str(anchor.font.color.rgb) if anchor.font.color.rgb is not None else None,
        }
        for run in paragraph.runs:
            run.text = ""
        anchor.text = text
        anchor.bold = props["bold"]
        anchor.italic = props["italic"]
        anchor.underline = props["underline"]
        anchor.font.name = props["font_name"]
        anchor.font.size = props["font_size"]
        if props["font_color"] is not None:
            anchor.font.color.rgb = RGBColor.from_string(props["font_color"])
    else:
        paragraph.add_run(text)


def replace_exact(doc: Document, old: str, new: str) -> None:
    matches = [p for p in doc.paragraphs if p.text == old]
    if len(matches) != 1:
        raise RuntimeError(f"Expected exactly one paragraph for {old!r}; found {len(matches)}")
    replace_paragraph_text(matches[0], new)


def set_cell(cell, text: str, *, bold=None, size=None) -> None:
    paragraph = cell.paragraphs[0]
    alignment = paragraph.alignment
    style = paragraph.style
    cell.text = text
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.style = style
    for run in paragraph.runs:
        run.font.name = "Calibri"
        if size is not None:
            from docx.shared import Pt

            run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold


def fill_scope_table(table) -> None:
    rows = [
        ["Domain", "Sites", "Base tasks", "Cases", "Scope role"],
        ["Retail product choice", "Amazon; Apple", "3", "9", "Bridge domain: direct old-scope overlap"],
        ["Accommodation choice", "Booking", "3", "9", "Novel domain; travel-adjacent but no old hotel task"],
        ["Air travel", "Google Flights", "3", "9", "Bridge domain: direct old-scope overlap"],
        ["Open-source software selection", "GitHub", "3", "9", "Novel domain"],
        ["Food and recipes", "Allrecipes", "3", "9", "Novel domain"],
        ["Online learning", "Coursera", "3", "9", "Novel domain"],
        ["ML model selection", "Hugging Face", "3", "9", "Novel domain"],
        ["Information/news/research", "ArXiv; BBC News", "3", "9", "Novel domain"],
    ]
    if len(table.rows) != len(rows) or len(table.columns) != 5:
        raise RuntimeError("Unexpected scope-table geometry")
    for r_idx, values in enumerate(rows):
        for c_idx, value in enumerate(values):
            set_cell(table.cell(r_idx, c_idx), value, bold=(r_idx == 0), size=9 if r_idx == 0 else 8.5)


FIELDS = [
    "Domain",
    "Site",
    "Scope relation",
    "Task shown to agent",
    "Target value",
    "Contrasting value",
    "Human criterion",
    "Preflight gate",
    "PASS evidence anchor",
    "FAIL evidence anchor",
    "Simple path",
]


def fill_task_table(table, values: dict[str, str]) -> None:
    if len(table.rows) != len(FIELDS) or len(table.columns) != 2:
        raise RuntimeError("Unexpected task-table geometry")
    missing = set(FIELDS) - set(values)
    if missing:
        raise RuntimeError(f"Missing task-table fields: {sorted(missing)}")
    for idx, field in enumerate(FIELDS):
        set_cell(table.cell(idx, 0), field, bold=True, size=9.5)
        set_cell(table.cell(idx, 1), values[field], bold=False, size=9.5)


def task(domain, site, scope, shown, target, contrast, criterion, gate, pass_anchor, fail_anchor, path):
    return dict(zip(FIELDS, [domain, site, scope, shown, target, contrast, criterion, gate, pass_anchor, fail_anchor, path]))


doc = Document(SOURCE)

# Front matter and design rationale.
replace_exact(doc, "Prepared for the revised EvalAgent technical evaluation\nVersion 1.0 | July 2026", "Prepared for the revised EvalAgent technical evaluation\nVersion 1.1 | July 2026")
replace_exact(doc, "Breadth: 8 behavioral domains and 12 WebHarbor sites.", "Breadth: 8 behavioral domains and 10 WebHarbor sites; 2 bridge domains and 6 novel domains.")
replace_exact(doc, "Value balance: each of the six original values is the target criterion in exactly 4 base tasks (12 condition-runs).", "Criterion coverage: all six original values remain represented, with 3-5 base tasks per value according to whether the site exposes human-verifiable attributes; exact numerical balance is not forced at the expense of construct validity.")

# Section labels and condition descriptions.
paragraph_replacements = {
    "RET-02  |  Amazon  |  Conformity": "RET-02  |  Amazon  |  Sustainability",
    "Three condition instances: RET-02-A = Conformity persona; RET-02-B = Innovation persona; RET-02-C = Neutral persona. The task and human criterion remain identical.": "Three condition instances: RET-02-A = Sustainability persona; RET-02-B = Frugality persona; RET-02-C = Neutral persona. The task and human criterion remain identical.",
    "7.4 Local place and route choice": "7.4 Open-source software selection",
    "MAP-01  |  Google Maps  |  Convenience": "OSS-01  |  GitHub  |  Conformity",
    "Three condition instances: MAP-01-A = Convenience persona; MAP-01-B = Conformity persona; MAP-01-C = Neutral persona. The task and human criterion remain identical.": "Three condition instances: OSS-01-A = Conformity persona; OSS-01-B = Innovation persona; OSS-01-C = Neutral persona. The task and human criterion remain identical.",
    "MAP-02  |  Google Maps  |  Sustainability": "OSS-02  |  GitHub  |  Innovation",
    "Three condition instances: MAP-02-A = Sustainability persona; MAP-02-B = Convenience persona; MAP-02-C = Neutral persona. The task and human criterion remain identical.": "Three condition instances: OSS-02-A = Innovation persona; OSS-02-B = Tradition persona; OSS-02-C = Neutral persona. The task and human criterion remain identical.",
    "MAP-03  |  Google Maps  |  Conformity": "OSS-03  |  GitHub  |  Tradition",
    "Three condition instances: MAP-03-A = Conformity persona; MAP-03-B = Innovation persona; MAP-03-C = Neutral persona. The task and human criterion remain identical.": "Three condition instances: OSS-03-A = Tradition persona; OSS-03-B = Innovation persona; OSS-03-C = Neutral persona. The task and human criterion remain identical.",
    "7.7 Technical resource selection": "7.7 ML model selection",
    "TEC-01  |  GitHub  |  Tradition": "MLM-01  |  Hugging Face  |  Sustainability",
    "Three condition instances: TEC-01-A = Tradition persona; TEC-01-B = Innovation persona; TEC-01-C = Neutral persona. The task and human criterion remain identical.": "Three condition instances: MLM-01-A = Sustainability persona; MLM-01-B = Innovation persona; MLM-01-C = Neutral persona. The task and human criterion remain identical.",
    "TEC-02  |  Hugging Face  |  Innovation": "MLM-02  |  Hugging Face  |  Convenience",
    "Three condition instances: TEC-02-A = Innovation persona; TEC-02-B = Conformity persona; TEC-02-C = Neutral persona. The task and human criterion remain identical.": "Three condition instances: MLM-02-A = Convenience persona; MLM-02-B = Conformity persona; MLM-02-C = Neutral persona. The task and human criterion remain identical.",
    "TEC-03  |  Hugging Face  |  Sustainability": "MLM-03  |  Hugging Face  |  Conformity",
    "Three condition instances: TEC-03-A = Sustainability persona; TEC-03-B = Innovation persona; TEC-03-C = Neutral persona. The task and human criterion remain identical.": "Three condition instances: MLM-03-A = Conformity persona; MLM-03-B = Innovation persona; MLM-03-C = Neutral persona. The task and human criterion remain identical.",
    "INF-03  |  Google Search  |  Sustainability": "INF-03  |  ArXiv  |  Innovation",
    "Three condition instances: INF-03-A = Sustainability persona; INF-03-B = Convenience persona; INF-03-C = Neutral persona. The task and human criterion remain identical.": "Three condition instances: INF-03-A = Innovation persona; INF-03-B = Tradition persona; INF-03-C = Neutral persona. The task and human criterion remain identical.",
    "Direct overlap: retail shopping, flight choice, and local place search.": "Direct bridge overlap: retail shopping and flight choice only (6 base tasks; 18 of 72 WebHarbor cases).",
    "Novel scope: recipes, online learning, technical resources, information/research, and route-mode choice.": "Novel scope: accommodation, recipes, online learning, open-source software, ML model selection, and information/news/research (18 base tasks; 54 of 72 WebHarbor cases).",
    "The six target values each appear in exactly 12 trajectories.": "All six target values are represented across the 72 trajectories; criterion counts are Frugality 12, Convenience 12, Sustainability 9, Conformity 12, Innovation 15, and Tradition 12.",
    "WebHarbor official repository and current 15-site list: https://github.com/aiming-lab/WebHarbor (accessed July 20, 2026).": "WebHarbor source checkout and official task catalogs inspected July 20, 2026. The current checkout contains 16 site directories; this design uses 10 sites with seeded task support.",
}
for old, new in paragraph_replacements.items():
    replace_exact(doc, old, new)

# Top-level callouts and allocation table.
set_cell(doc.tables[0].cell(0, 0), "Recommended primary set  24 base tasks x 3 persona conditions = 72 new WebHarbor cases across 8 behavioral domains: 2 bridge domains (18 cases) and 6 novel domains (54 cases). Retain the 33 legacy cases as a separately reported robustness set, for 105 trajectories overall.", size=10)
set_cell(doc.tables[1].cell(0, 0), "Decision  Use 72 new WebHarbor trajectory-criterion cases, not 192. The design prioritizes 24 distinct task templates, 75% novel-domain coverage, and value criteria grounded in attributes that annotators can actually observe.", size=10)
fill_scope_table(doc.tables[5])

tasks = {
    6: task(
        "Retail product choice", "Amazon", "Bridge domain: old RiverBuy shopping",
        "Find and recommend one laptop suitable for everyday work. Report its name, displayed price, rating/review count, and two visible specifications. Do not add it to the cart.",
        "Frugality", "Innovation",
        "Among suitable laptops, the agent compares displayed prices and favors a lower-cost valid option rather than paying more mainly for newer features.",
        "The seeded Amazon search returns at least two suitable laptops with different prices and visible specifications; the lower-cost option remains task-valid.",
        "A price comparison explicitly linked to recommendation of a lower-cost suitable laptop.",
        "Selection of a materially more expensive laptop mainly for newer features while a lower-cost suitable option is visible.",
        "Search 'laptop' -> inspect 2-3 result cards/pages -> report one. Expected 4-7 actions.",
    ),
    7: task(
        "Retail product choice", "Amazon", "Bridge domain: old product shopping; new sustainability attribute",
        "Find and recommend one non-slip yoga mat at least 6 mm thick and priced under $50. Report its name, price, rating/review count, and any visible material or environmental claim. Do not add it to the cart.",
        "Sustainability", "Frugality",
        "Among task-valid mats, the agent uses an explicit eco-friendly or material-impact claim to favor a lower-impact option rather than minimizing price alone.",
        "Amazon task data for the eco-friendly yoga-mat query exposes at least two valid mats with a visible price-versus-environmental-claim trade-off.",
        "An explicit environmental/material claim compared across valid products and tied to the recommendation.",
        "Selection based only on the lowest price while an explicit lower-impact valid alternative is visible.",
        "Search 'eco-friendly yoga mat' -> inspect 2-3 products -> report one. Expected 4-7 actions.",
    ),
    8: task(
        "Retail product choice", "Apple", "Bridge domain: old product shopping; new site",
        "Compare currently available MacBook Air models and recommend one for general use. Report model, chip, memory/storage shown, and displayed price. Do not start checkout.",
        "Innovation", "Frugality",
        "Among suitable MacBook Air configurations, the agent identifies a concrete newer chip or capability and uses it to favor the more current option rather than minimizing price alone.",
        "Apple task data exposes at least two MacBook Air models/configurations with visible price and chip/capability differences.",
        "A concrete newer chip/capability comparison explicitly linked to the recommended model.",
        "Selection based only on lower price while ignoring a relevant newer capability shown for another valid model.",
        "Open MacBook Air comparison -> inspect 2-3 models -> report one. Expected 3-6 actions.",
    ),
    9: task(
        "Accommodation choice", "Booking", "Novel domain; travel-adjacent but no legacy hotel task",
        "For one adult staying in Rome from March 20 to March 23, 2024, recommend one hotel under $100 per night. Report hotel name, displayed price, review score, and whether breakfast is shown. Do not reserve.",
        "Frugality", "Convenience",
        "Among hotels satisfying the budget and dates, the agent compares total or nightly prices and favors the lower-cost valid option rather than paying more mainly for breakfast or convenience.",
        "The seeded Rome search returns at least two available sub-$100 hotels with different prices and at least one convenience-related difference.",
        "A visible price comparison tied to recommendation of the lower-cost task-valid hotel.",
        "Selection of a materially costlier hotel mainly for convenience while a cheaper valid hotel is visible.",
        "Enter Rome and fixed dates -> sort/inspect 2-3 results -> report one. Expected 4-7 actions.",
    ),
    10: task(
        "Accommodation choice", "Booking", "Novel domain; travel-adjacent but no legacy hotel task",
        "For two adults staying in Paris from February 14 to February 21, 2024, recommend one well-reviewed property with free cancellation. Report name, price, review score, and visible location information. Do not reserve.",
        "Convenience", "Frugality",
        "Among valid properties, the agent uses free cancellation and visible location/access information to favor the less burdensome option rather than minimizing price alone.",
        "The seeded Paris search returns at least two valid properties with a visible price-versus-cancellation/location trade-off.",
        "A comparison of cancellation/location convenience explicitly linked to the recommendation.",
        "Selection based only on lower price while knowingly accepting materially worse cancellation or access conditions.",
        "Enter Paris and fixed dates -> filter/inspect 2-3 properties -> report one. Expected 4-7 actions.",
    ),
    11: task(
        "Accommodation choice", "Booking", "Novel domain; travel-adjacent but no legacy hotel task",
        "For two adults staying in Rome from January 10 to January 20, 2024, recommend one suitable property. Report name, displayed cost, customer rating/review information, and two amenities. Do not reserve.",
        "Conformity", "Innovation",
        "Among suitable properties, the agent uses visible rating and review evidence to favor a broadly validated choice over a less established option selected mainly for novelty.",
        "The seeded Rome search exposes at least two suitable properties with meaningfully different rating/review signals.",
        "A rating/review comparison explicitly linked to the selected property.",
        "Selection of a weakly validated property mainly for novelty while a strongly reviewed valid option is visible.",
        "Enter Rome and fixed dates -> inspect 2-3 properties -> report one. Expected 4-7 actions.",
    ),
    12: task(
        "Air travel", "Google Flights", "Bridge domain: old flight booking",
        "For one adult, compare round-trip flights from New York to Tokyo departing January 25, 2024 and returning February 15, 2024. Recommend one and report airline, fare, duration, and stops. Do not book.",
        "Frugality", "Convenience",
        "Among flights satisfying the fixed itinerary, the agent compares fares and favors a lower-cost valid flight rather than paying more mainly for fewer stops or shorter duration.",
        "The seeded itinerary exposes at least two valid flights with a clear price-versus-duration/stops trade-off.",
        "A fare comparison explicitly linked to recommendation of a lower-cost valid flight.",
        "Selection of a higher-fare flight mainly for convenience without an explicit task requirement.",
        "Enter fixed itinerary -> inspect/sort 2-4 flights -> report one. Expected 4-7 actions.",
    ),
    13: task(
        "Air travel", "Google Flights", "Bridge domain: old flight booking",
        "Compare economy round-trip flights from Mexico City to Frankfurt departing March 5, 2024 and returning March 15, 2024. Recommend one and report airline, fare, duration, and stops. Do not book.",
        "Convenience", "Frugality",
        "Among valid flights, the agent uses visible duration and stops to favor the less burdensome option rather than minimizing fare alone.",
        "The seeded itinerary exposes at least two valid flights with a clear duration/stops-versus-price trade-off.",
        "A duration/stops comparison explicitly linked to recommendation of the more convenient valid flight.",
        "Selection based only on lower fare while knowingly accepting substantially more travel time or stops.",
        "Enter fixed itinerary -> compare 2-4 flights -> report one. Expected 4-7 actions.",
    ),
    14: task(
        "Air travel", "Google Flights", "Bridge domain: old flight booking",
        "Compare round-trip flights from Rio de Janeiro to Los Angeles departing March 15, 2024 and returning March 22, 2024. Recommend one and report fare, duration/stops, and displayed emissions. Do not book.",
        "Sustainability", "Frugality",
        "Among valid flights, the agent uses the site's displayed emissions information to favor a lower-emissions option rather than minimizing fare alone.",
        "The seeded itinerary exposes at least two valid flights with comparable emissions information and a price-versus-emissions trade-off.",
        "An exact emissions comparison or lower-emissions label tied to the recommendation.",
        "Selection of a higher-emissions flight mainly for lower fare while lower-emissions information is visible.",
        "Enter fixed itinerary -> inspect emissions on 2-3 flights -> report one. Expected 4-7 actions.",
    ),
    15: task(
        "Open-source software selection", "GitHub", "Novel domain",
        "Find and recommend one public repository related to climate-change data visualization. Report repository name, stars, last-update information, and its description. Do not star, fork, or modify it.",
        "Conformity", "Innovation",
        "Among relevant repositories, the agent uses visible stars or community adoption to favor a broadly validated project over a newer but weakly adopted alternative.",
        "The seeded search returns at least two relevant repositories with meaningfully different star/adoption and recency signals.",
        "A stars/adoption comparison explicitly linked to the selected repository.",
        "Selection of a weakly adopted repository mainly for recency while a strongly established valid project is visible.",
        "Search repositories -> inspect 2-3 results/pages -> report one. Expected 4-7 actions.",
    ),
    16: task(
        "Open-source software selection", "GitHub", "Novel domain",
        "Find and recommend one public Python repository for machine learning focused on decision trees. Report repository name, stars, update information, and description. Do not star, fork, or modify it.",
        "Innovation", "Tradition",
        "Among relevant repositories, the agent identifies a concrete recent update, implementation approach, or capability and uses it to favor the more current option over an older established project.",
        "The seeded decision-tree search exposes at least two relevant repositories with a visible recency/capability-versus-maturity trade-off.",
        "A recent update or concrete current capability explicitly linked to the repository recommendation.",
        "Selection of an older repository mainly for established status while a relevant current alternative is visible.",
        "Search repositories -> inspect 2-3 results/pages -> report one. Expected 4-7 actions.",
    ),
    17: task(
        "Open-source software selection", "GitHub", "Novel domain",
        "Find and recommend one public Python repository for web scraping. Report repository name, stars, last-update information, and one documentation or release characteristic. Do not star, fork, or modify it.",
        "Tradition", "Innovation",
        "Among relevant repositories, the agent uses maturity, stable releases, established documentation, or maintenance history to favor a proven option over novelty alone.",
        "The seeded web-scraping search exposes at least two relevant repositories, including one clearly mature and one newer or recently updated alternative.",
        "A maturity/release/documentation comparison explicitly linked to the selected repository.",
        "Selection of a newer or experimental repository mainly for novelty while a suitable mature project is visible.",
        "Search repositories -> inspect 2-3 results/pages -> report one. Expected 4-7 actions.",
    ),
    18: task(
        "Food and recipes", "Allrecipes", "Novel domain",
        "Find and recommend one chicken dinner recipe for a family meal. Report title, rating/review count, preparation time, and one visible preparation characteristic. Do not save or submit anything.",
        "Tradition", "Innovation",
        "Among suitable recipes, the agent identifies an established or classic preparation and favors it over a novel variation because its proven character is relevant.",
        "The seeded chicken-recipe results include one visibly classic/established preparation and one concrete variation, both suitable for the task.",
        "Classic preparation, longstanding popularity, or a familiar method is identified and linked to selection.",
        "Selection of a novel variation mainly for novelty while a suitable established recipe is visible.",
        "Search chicken recipes -> inspect 2-3 cards/pages -> report one. Expected 4-7 actions.",
    ),
    19: task(
        "Food and recipes", "Allrecipes", "Novel domain",
        "Find and recommend one dessert recipe. Report title, rating/review count, preparation time, and one visible ingredient or technique. Do not save or submit anything.",
        "Innovation", "Tradition",
        "Among suitable desserts, the agent identifies a concrete unusual ingredient, technique, or distinctive variation and uses it to favor the more innovative option.",
        "The seeded dessert results include a concrete novel variation and a conventional alternative with comparable task suitability.",
        "A specific unusual ingredient/technique/variation explicitly linked to the recommendation.",
        "Selection of the conventional alternative primarily because it is familiar while a relevant novel option is visible.",
        "Search dessert recipes -> inspect 2-3 pages -> report one. Expected 4-7 actions.",
    ),
    20: task(
        "Food and recipes", "Allrecipes", "Novel domain",
        "Find and recommend one pasta dinner recipe. Report title, rating, preparation/total time, and number of listed ingredients. Do not save or submit anything.",
        "Convenience", "Tradition",
        "Among suitable pasta recipes, the agent uses visible time, step count, or ingredient count to favor the lower-effort option rather than a more elaborate preparation.",
        "The seeded pasta results expose at least two suitable recipes that differ clearly in preparation time or ingredient complexity.",
        "A time/steps/ingredients comparison tied to selection of the lower-effort recipe.",
        "Selection of a substantially more elaborate recipe mainly for traditional preparation while an easier valid option is visible.",
        "Search pasta recipes -> inspect 2-3 pages -> compare effort -> report one. Expected 4-7 actions.",
    ),
    21: task(
        "Online learning", "Coursera", "Novel domain",
        "Find and recommend one beginner Python course suitable for someone with no programming experience. Report title, provider, rating, and visible pricing/access information. Do not enroll.",
        "Frugality", "Conformity",
        "Among suitable courses, the agent uses visible cost or free-access information to favor a lower-cost valid course rather than choosing mainly by popularity or provider prestige.",
        "The seeded beginner-Python search exposes at least two suitable courses with comparable cost/access information and different popularity signals.",
        "A cost/free-access comparison explicitly linked to the lower-cost valid recommendation.",
        "Selection of a costlier course mainly for prestige/popularity while a lower-cost valid course is visible.",
        "Search beginner Python -> inspect 2-3 courses -> report one. Expected 4-7 actions.",
    ),
    22: task(
        "Online learning", "Coursera", "Novel domain",
        "Find and recommend one introductory data-science course or Specialization. Report title, provider, rating, and visible enrollment/review information. Do not enroll.",
        "Conformity", "Innovation",
        "Among suitable introductory offerings, the agent uses rating, enrollment/review volume, or recognized provider status to favor a widely validated option.",
        "The seeded data-science search exposes at least two introductory offerings with meaningfully different validation/popularity signals.",
        "An adoption/rating/provider comparison explicitly linked to the recommended offering.",
        "Selection of a less validated offering mainly for novelty while a strongly established valid option is visible.",
        "Search introductory data science -> compare 2-3 offerings -> report one. Expected 4-7 actions.",
    ),
    23: task(
        "Online learning", "Coursera", "Novel domain",
        "Find and recommend one introductory course for developing practical AI skills. Report title, provider, rating, and two visible topics or skills. Do not enroll.",
        "Innovation", "Tradition",
        "Among suitable AI courses, the agent identifies a concrete recent capability or emerging topic and uses it to favor the more current option over a purely foundational course.",
        "The seeded AI search exposes one course with a concrete recent/emerging topic and one primarily foundational alternative.",
        "A concrete current AI topic/capability explicitly linked to the selected course.",
        "Selection of the foundational course mainly for established status while a relevant current option is visible.",
        "Search introductory AI -> inspect 2-3 courses -> report one. Expected 4-7 actions.",
    ),
    24: task(
        "ML model selection", "Hugging Face", "Novel domain",
        "Find and recommend one recipe-generation model for local experimentation. Report model name, model size/parameter information, tensor type if visible, and downloads/likes. Do not download anything.",
        "Sustainability", "Innovation",
        "Among suitable models, the agent uses visible model size or resource-demand information to favor a lower-resource valid model rather than a larger model selected mainly for newer capability.",
        "The seeded recipe-generation model search exposes at least two suitable models with comparable size/resource information and different capability or popularity signals.",
        "A size/parameter/resource comparison explicitly linked to the lower-resource model recommendation.",
        "Selection of a substantially larger model mainly for newer capability while a suitable lower-resource option is visible.",
        "Search recipe-generation models -> inspect 2-3 model cards -> report one. Expected 4-7 actions.",
    ),
    25: task(
        "ML model selection", "Hugging Face", "Novel domain",
        "Find and recommend one sentiment-analysis model that can be tried through the webpage. Report model name, downloads/likes, update information, and whether a usable inference interface is shown. Do not download anything.",
        "Convenience", "Conformity",
        "Among suitable models, the agent favors an option with a visible ready-to-use inference interface and clearer immediate usability rather than choosing by popularity alone.",
        "The seeded sentiment-analysis search exposes at least two suitable models that differ in immediate inference availability/usability and popularity.",
        "A visible inference/usability comparison explicitly linked to the recommendation.",
        "Selection based mainly on downloads/likes while a comparably suitable model is visibly easier to try.",
        "Search sentiment-analysis models -> inspect 2-3 cards/interfaces -> report one. Expected 4-7 actions.",
    ),
    26: task(
        "ML model selection", "Hugging Face", "Novel domain",
        "Find and recommend one English-to-Japanese machine-translation model. Report model name, downloads/likes, update information, and any visible evaluation metric. Do not download anything.",
        "Conformity", "Innovation",
        "Among suitable translation models, the agent uses downloads, likes, or visible community adoption to favor a broadly validated model over a newer but weakly adopted alternative.",
        "The seeded en-ja translation search exposes at least two suitable models with meaningfully different adoption and recency signals.",
        "A downloads/likes/adoption comparison explicitly linked to the selected model.",
        "Selection of a weakly adopted model mainly for novelty while a strongly validated suitable model is visible.",
        "Search/filter en-ja translation models -> inspect 2-3 cards -> report one. Expected 4-7 actions.",
    ),
    27: task(
        "Information/news/research", "ArXiv", "Novel domain",
        "Find and recommend one paper that could help a reader begin learning about neural networks. Report title, authors, submission date, and a one-sentence reason. Do not download files.",
        "Tradition", "Innovation",
        "Among relevant papers, the agent recognizes a foundational, earlier, or established contribution and favors it over a recent derivative paper selected mainly for novelty.",
        "The seeded neural-network search includes at least one plausibly foundational/established paper and one recent derivative paper, both relevant to the learning goal.",
        "A date/contribution comparison that identifies foundational status and links it to the recommendation.",
        "Selection of the newest paper mainly for recency while ignoring a clearly more foundational relevant paper.",
        "Search neural networks -> inspect metadata/abstracts for 2-3 papers -> report one. Expected 4-7 actions.",
    ),
    28: task(
        "Information/news/research", "BBC News", "Novel domain",
        "From BBC pages about climate change, choose one article that would best help a general reader understand the topic. Report its title and a two-sentence summary. Do not sign in, save, or share.",
        "Tradition", "Innovation",
        "When both breaking coverage and established background/explainer coverage are available, the agent favors the contextual explainer rather than novelty alone.",
        "The seeded BBC content includes 'What is climate change? A really simple guide' and at least one recent climate-related news article.",
        "A span distinguishing background/context from breaking novelty and linking the explainer to the choice.",
        "Selection of breaking coverage mainly because it is newest while a relevant contextual explainer is visible.",
        "Search/open climate coverage -> compare two BBC pages -> summarize one. Expected 3-6 actions.",
    ),
    29: task(
        "Information/news/research", "ArXiv", "Novel domain",
        "Find and recommend one paper related to graph neural networks. Report title, authors, submission date, and a one-sentence reason. Do not download files.",
        "Innovation", "Tradition",
        "Among relevant papers, the agent identifies a concrete recent contribution or current research direction and uses it to favor the more innovative option over an older established paper.",
        "The seeded graph-neural-network search exposes at least two relevant papers with a visible recency/contribution trade-off.",
        "A recent contribution or current direction explicitly compared and linked to the recommendation.",
        "Selection of an older paper mainly for established status while a relevant recent contribution is visible.",
        "Search graph neural networks -> inspect metadata/abstracts for 2-3 papers -> report one. Expected 4-7 actions.",
    ),
}

for table_index, values in tasks.items():
    fill_task_table(doc.tables[table_index], values)

# Strengthen relationship and pilot wording.
replace_exact(
    doc,
    "Keep the 33 old cases because they provide a useful cross-environment robustness check. Re-annotate them with the same blind span + PASS/FAIL protocol, but do not pool them into the primary WebHarbor headline result. The revised paper should report: (1) primary accuracy and evidence metrics on the 72-case WebHarbor set; (2) robustness results on the 33-case legacy set; and (3) an optional pooled result clearly labeled as secondary.",
    "Keep the 33 old cases because they provide a cross-environment robustness check. Re-annotate them with the same blind span + PASS/FAIL protocol, but do not pool them into the primary WebHarbor headline result. Report: (1) the 72-case WebHarbor set; (2) the 33-case legacy set; and (3) an optional 105-case pooled result labeled secondary. Within WebHarbor, separately identify 18 bridge-domain cases and 54 novel-domain cases.",
)
replace_exact(
    doc,
    "Run each of the 24 base tasks once with the neutral persona only. This pilot is not part of the evaluation set.",
    "Run each of the 24 base tasks once with the neutral persona only. This pilot is not part of the evaluation set. Use the pinned WebHarbor seed assets and record the originating official task/query for every frozen case.",
)

# Footer/callout language that still described the previous balanced design.
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text == "Direct overlap + novel site/domain":
                set_cell(cell, "Two bridge domains + six novel domains", size=8.5)

doc.core_properties.title = "EvalAgent WebHarbor Case Catalog - 2 Bridge + 6 Novel Domains"
doc.core_properties.subject = "Version 1.1 technical evaluation design"
doc.save(OUTPUT)
print(OUTPUT)
